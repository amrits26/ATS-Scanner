"""
docx_generator.py

Generate Microsoft Word DOCX files from structured resume JSON with tracked changes.
Uses python-docx library.

Usage:
    from backend.services.docx_generator import generate_resume_docx
    docx_bytes = await generate_resume_docx(resume_json, tracked_changes)
"""

import io
import logging
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("python-docx not installed. Run: pip install python-docx")

logger = logging.getLogger(__name__)


def _add_tracked_change_highlight(paragraph, is_addition=False, is_deletion=False):
    """
    Add tracked change highlighting to a paragraph.
    Green for additions, red for deletions.
    """
    try:
        pPr = paragraph._element.get_or_add_pPr()
        
        if is_addition:
            # Green background for added text
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "90EE90")  # Light green
            pPr.append(shd)
        elif is_deletion:
            # Red background for deleted text (with strikethrough)
            for run in paragraph.runs:
                run.font.strike = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    except Exception as e:
        logger.warning(f"Could not apply tracked change highlighting: {e}")


async def generate_resume_docx(
    resume_json: Dict[str, Any],
    tracked_changes: Optional[List[Dict[str, str]]] = None,
) -> bytes:
    """
    Generate a professional resume DOCX from structured JSON.
    
    **Input:**
    ```json
    {
        "summary": "...",
        "skills": ["skill1", "skill2", ...],
        "experience": [
            {"company": "...", "title": "...", "dates": "...", "bullets": ["...", "..."]}
        ],
        "education": [
            {"degree": "...", "school": "...", "year": "..."}
        ],
        "projects": [
            {"name": "...", "description": "..."}
        ]
    }
    ```
    
    **Returns:**
    bytes: DOCX file content (can be written to S3)
    """
    
    try:
        doc = Document()
        
        # Set default style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        
        # ===== SUMMARY =====
        if resume_json.get("summary"):
            title = doc.add_paragraph()
            title_run = title.add_run("PROFESSIONAL SUMMARY")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            summary_p = doc.add_paragraph(resume_json["summary"])
            summary_p.paragraph_format.left_indent = Inches(0.25)
            doc.add_paragraph()  # Spacing
        
        # ===== SKILLS =====
        if resume_json.get("skills"):
            title = doc.add_paragraph()
            title_run = title.add_run("SKILLS")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            skills_text = ", ".join(resume_json["skills"][:15])
            skills_p = doc.add_paragraph(skills_text)
            skills_p.paragraph_format.left_indent = Inches(0.25)
            doc.add_paragraph()  # Spacing
        
        # ===== EXPERIENCE =====
        if resume_json.get("experience"):
            title = doc.add_paragraph()
            title_run = title.add_run("PROFESSIONAL EXPERIENCE")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            for job in resume_json["experience"]:
                # Company & Title on same line
                header = doc.add_paragraph()
                header.paragraph_format.left_indent = Inches(0.25)
                
                company_run = header.add_run(f"{job.get('company', 'Company')}")
                company_run.bold = True
                
                header.add_run(" | ")
                
                title_run = header.add_run(f"{job.get('title', 'Title')}")
                title_run.italic = True
                
                # Dates
                dates_p = doc.add_paragraph(f"Dates: {job.get('dates', 'Start Date - End Date')}")
                dates_p.paragraph_format.left_indent = Inches(0.5)
                dates_p.runs[0].font.size = Pt(10)
                dates_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                
                # Bullets
                for bullet in job.get("bullets", []):
                    bullet_p = doc.add_paragraph(bullet, style="List Bullet")
                    bullet_p.paragraph_format.left_indent = Inches(0.5)
                    bullet_p.paragraph_format.first_line_indent = Inches(-0.25)
                
                doc.add_paragraph()  # Spacing between jobs
        
        # ===== EDUCATION =====
        if resume_json.get("education"):
            title = doc.add_paragraph()
            title_run = title.add_run("EDUCATION")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            for edu in resume_json["education"]:
                edu_p = doc.add_paragraph()
                edu_p.paragraph_format.left_indent = Inches(0.25)
                
                degree_run = edu_p.add_run(f"{edu.get('degree', 'Degree')}")
                degree_run.bold = True
                
                edu_p.add_run(f" | {edu.get('school', 'School')}")
                
                if edu.get("year"):
                    year_run = edu_p.add_run(f" ({edu['year']})")
                    year_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
        
        # ===== PROJECTS (if any) =====
        if resume_json.get("projects"):
            title = doc.add_paragraph()
            title_run = title.add_run("NOTABLE PROJECTS")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            for project in resume_json["projects"]:
                proj_p = doc.add_paragraph()
                proj_p.paragraph_format.left_indent = Inches(0.25)
                
                name_run = proj_p.add_run(f"{project.get('name', 'Project')}: ")
                name_run.bold = True
                
                proj_p.add_run(project.get("description", "Project description"))
            
            doc.add_paragraph()  # Spacing
        
        # ===== TRACKED CHANGES SUMMARY (if provided) =====
        if tracked_changes:
            doc.add_paragraph()  # Spacing
            
            changes_title = doc.add_paragraph()
            changes_title_run = changes_title.add_run("CHANGES MADE FOR THIS JOB")
            changes_title_run.bold = True
            changes_title_run.font.size = Pt(11)
            changes_title_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            
            # Categorize changes
            additions = [c for c in tracked_changes if c.get("type") == "added"]
            modifications = [c for c in tracked_changes if c.get("type") == "modified"]
            removals = [c for c in tracked_changes if c.get("type") == "removed"]
            
            # Show top 5 key changes
            if modifications:
                mod_title = doc.add_paragraph()
                mod_title_run = mod_title.add_run("Key Modifications:")
                mod_title_run.italic = True
                mod_title.paragraph_format.left_indent = Inches(0.25)
                
                for change in modifications[:5]:
                    mod_p = doc.add_paragraph(style="List Bullet")
                    mod_p.paragraph_format.left_indent = Inches(0.5)
                    mod_p.add_run(f"'{change.get('original', '')[:50]}' → '{change.get('rewritten', '')[:50]}'")
            
            if additions:
                add_title = doc.add_paragraph()
                add_title_run = add_title.add_run("Key Additions:")
                add_title_run.italic = True
                add_title.paragraph_format.left_indent = Inches(0.25)
                
                for change in additions[:3]:
                    add_p = doc.add_paragraph(style="List Bullet")
                    add_p.paragraph_format.left_indent = Inches(0.5)
                    add_p.add_run(change.get("rewritten", "")[:80])
                    _add_tracked_change_highlight(add_p, is_addition=True)
        
        # Convert to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"[DOCX] Generated resume DOCX ({len(buffer.getvalue())} bytes)")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"[DOCX] Generation error: {e}")
        raise


async def generate_sample_docx_for_testing() -> bytes:
    """Generate a sample DOCX for testing."""
    sample_json = {
        "summary": "Results-driven Software Engineer with 5+ years of experience building scalable backend systems.",
        "skills": [
            "Python", "Django", "FastAPI", "PostgreSQL", "Docker", "Kubernetes",
            "AWS", "Redis", "GraphQL", "React", "TypeScript"
        ],
        "experience": [
            {
                "company": "TechCorp Inc",
                "title": "Senior Backend Engineer",
                "dates": "Jan 2023 - Present",
                "bullets": [
                    "Led redesign of core API using FastAPI, reducing response latency by 40%",
                    "Managed PostgreSQL optimization resulting in 3M+ queries/day at <100ms p95"
                ]
            }
        ],
        "education": [
            {
                "degree": "BS Computer Science",
                "school": "University of California",
                "year": "2018"
            }
        ],
        "projects": []
    }
    
    return await generate_resume_docx(sample_json, tracked_changes=[])
