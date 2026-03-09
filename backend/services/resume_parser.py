"""Resume text extraction from PDF and DOCX."""

import io
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document

from ..utils.text_cleaner import clean_extracted_text


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes. Handles multi-page and layout issues."""
    if not content:
        return ""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            parts: list[str] = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            raw = "\n".join(parts) if parts else ""
            return clean_extracted_text(raw)
    except Exception:
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    if not content:
        return ""
    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        raw = "\n".join(parts) if parts else ""
        return clean_extracted_text(raw)
    except Exception:
        return ""


def extract_resume_text(content: bytes, filename: str) -> str:
    """
    Dispatch by extension. Returns cleaned text or empty string on failure.
    """
    path = Path(filename or "")
    suffix = (path.suffix or "").lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(content)
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(content)
    return ""
