"""Generate downloadable DOCX from optimized resume text."""

import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _split_into_sections(text: str) -> list[tuple[str, str]]:
    """
    Parse plain text into (heading, content) pairs.
    Looks for ALL CAPS or common section titles.
    """
    sections: list[tuple[str, str]] = []
    if not text:
        return sections
    # Normalize line breaks
    text = re.sub(r"\r\n", "\n", text).strip()
    # Section headers: line that is mostly caps or known titles
    known = {"summary", "experience", "education", "skills", "projects", "certifications", "contact"}
    lines = text.split("\n")
    current_heading = "Resume"
    current_body: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_body:
                sections.append((current_heading, "\n".join(current_body)))
                current_body = []
            continue
        # Check if this line looks like a section header (all caps or known title)
        is_header = (
            stripped.isupper() and len(stripped) < 50
            or stripped.lower().rstrip("s") in known
            or stripped.lower() in known
        )
        if is_header and len(stripped) < 60:
            if current_body:
                sections.append((current_heading, "\n".join(current_body)))
                current_body = []
            current_heading = stripped
        else:
            current_body.append(stripped)
    if current_body:
        sections.append((current_heading, "\n".join(current_body)))
    return sections


def generate_docx(optimized_resume: str) -> bytes:
    """
    Build a DOCX with clear headings and body. Returns file bytes.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    sections = _split_into_sections(optimized_resume or "")
    if not sections:
        doc.add_paragraph(optimized_resume or "No content.")
    for heading, body in sections:
        p_heading = doc.add_paragraph()
        run = p_heading.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        for block in body.split("\n"):
            block = block.strip()
            if block:
                style = "List Bullet" if block.startswith(("•", "-")) else "Normal"
                doc.add_paragraph(block, style=style)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
